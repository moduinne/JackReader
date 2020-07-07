from docx import Document
#returns all the paragraphs in the document by order
def obtainText(docFileName):
    document = Document(docFileName)
    finalText = []
    for line in document.paragraphs:
        finalText.append(line.text)
    return '\n'.join(finalText)

#returns all the tables in the document by order
def obtainTables(docFileName):
    document = Document(docFileName)
    finalTables = []
    for table in document.tables:
        finalTables.append(table)
    return finalTables

#returns a give tables rowList as a dictionaries
def returnTableRowList(table):
    data =  []
    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys,text))
        data.append(row_data)
    return data

#TODO
def returnColHeadings(data):
    return data[0]

def returnNumColDictionary(table):
    heads = []
    td = {}
    col = 0
    for i in table:
        heads.append(i)
    while col < len(table) :
        td[col] = heads[col]
        col += 1
    return td

print(obtainTables("Mock.docx")[0])
q = True

def runProram():
    inp = input("path: ")
    numb = 0
    if inp != "quit":
        print(len(obtainTables(inp)))
        while numb < 10:#len(obtainTables(inp)):
            table = returnTableRowList(obtainTables(inp)[numb])#[0]
            print(returnNumColDictionary(table))
            numb += 1
    if inp == "quit":
        exit()

while q:
    runProram()
    

    










    


#print(table)

#print(obtainText2("test.docx")[0])




###################################################################################
####################THIS IS THE DOCX2TXT VERSION###################################
###################################################################################

#import docx2txt
#my_text = docx2txt.process("test.docx")
#f = open("my_text_file.txt", "w")
#f.write(my_text)
#f.close()

#f = open("my_text_file.txt","r")
#line = f.readline()
#lineNum = 0
#while line != "":
#    print(str(lineNum) + ": " + line.strip())
#    line = f.readline()
#    lineNum += 1

    

